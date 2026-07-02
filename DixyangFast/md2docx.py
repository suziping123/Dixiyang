from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def md_to_docx(md_path, docx_path):
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        
        if line.startswith('# '):
            heading = doc.add_heading(level=1)
            run = heading.add_run(line[2:])
            run.bold = True
            run.font.size = Pt(16)
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        elif line.startswith('## '):
            heading = doc.add_heading(level=2)
            run = heading.add_run(line[3:])
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
        elif line.startswith('### '):
            heading = doc.add_heading(level=3)
            run = heading.add_run(line[4:])
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
        elif line.startswith('#### '):
            p = doc.add_paragraph()
            run = p.add_run(line[5:])
            run.bold = True
            run.font.size = Pt(11)
        
        elif line.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                p.paragraph_format.left_indent = Cm(1)
        
        elif line.startswith('|') and '|---' in line:
            i += 1
            while i < len(lines) and lines[i].startswith('|'):
                row = lines[i]
                cells = [c.strip() for c in row.split('|') if c.strip()]
                if len(cells) >= 3:
                    p = doc.add_paragraph()
                    p.add_run(f"{cells[0]}  {cells[1]}  {cells[2]}")
                i += 1
            continue
        
        elif line.startswith('|'):
            cells = [c.strip() for c in line.split('|') if c.strip()]
            if len(cells) >= 3:
                p = doc.add_paragraph()
                p.add_run(f"{cells[0]}  {cells[1]}  {cells[2]}")
        
        elif line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            text = line[2:]
            parts = []
            j = 0
            while j < len(text):
                if text[j:j+2] == '**':
                    j += 2
                    end = text.find('**', j)
                    if end != -1:
                        bold_text = text[j:end]
                        parts.append(('bold', bold_text))
                        j = end + 2
                    else:
                        parts.append(('text', text[j:]))
                        break
                elif text[j:j+1] == '`':
                    j += 1
                    end = text.find('`', j)
                    if end != -1:
                        code_text = text[j:end]
                        parts.append(('code', code_text))
                        j = end + 1
                    else:
                        parts.append(('text', text[j:]))
                        break
                else:
                    parts.append(('text', text[j]))
                    j += 1
            
            for part_type, part_text in parts:
                run = p.add_run(part_text)
                if part_type == 'bold':
                    run.bold = True
                elif part_type == 'code':
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0, 128, 0)
        
        elif line.strip() == '':
            doc.add_paragraph()
        
        else:
            p = doc.add_paragraph()
            text = line
            
            parts = []
            j = 0
            while j < len(text):
                if text[j:j+2] == '**':
                    j += 2
                    end = text.find('**', j)
                    if end != -1:
                        bold_text = text[j:end]
                        parts.append(('bold', bold_text))
                        j = end + 2
                    else:
                        parts.append(('text', text[j:]))
                        break
                elif text[j:j+1] == '`':
                    j += 1
                    end = text.find('`', j)
                    if end != -1:
                        code_text = text[j:end]
                        parts.append(('code', code_text))
                        j = end + 1
                    else:
                        parts.append(('text', text[j:]))
                        break
                else:
                    parts.append(('text', text[j]))
                    j += 1
            
            for part_type, part_text in parts:
                run = p.add_run(part_text)
                if part_type == 'bold':
                    run.bold = True
                elif part_type == 'code':
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0, 128, 0)
        
        i += 1
    
    doc.save(docx_path)
    print(f'转换完成: {md_path} -> {docx_path}')

if __name__ == '__main__':
    md_to_docx('daily-log/README.md', 'daily-log/README.docx')